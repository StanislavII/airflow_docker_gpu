#!/usr/bin/env python
# coding: utf-8

# In[2]:


import json
import posixpath
import numpy as np
import pandas as pd
from tqdm import tqdm
from datetime import date
import torch
import transformers
import torch.nn as nn
from datasets import load_dataset, Dataset
from transformers import AutoTokenizer, AutoModelWithLMHead, AutoModel
from datasets import load_dataset

import sys, getopt, os
import re
import docx



###


def preprocessor(text):
    text = re.sub('<[^>]*>', '', text)
    emoticons = re.findall('(?::|;|=)(?:-)?(?:\)|\(|D|P) ', text)
    text = (re.sub('[\W]+', ' ',  text.lower())+' '.join(emoticons).replace('-', ''))
    text = text.replace('n', '').replace('xa0', '')
    return text
    

def getText(filename, **kwargs):
    chunk_size_start = kwargs.get('start', None)
    chunk_size_end = kwargs.get('end', None)
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs[chunk_size_start:chunk_size_end]:
        fullText.append(para.text)
    return ' '.join(fullText)  
    
def search(list, platform):
    for i in range(len(list)):
        if list[i] == platform:
            return True
    return False
###
path = 'modules/docs2'
path_new = 'modules/docs3'

for root, dirs, files in os.walk(path):
    break

for root_new, dirs_new, files_new in os.walk(path_new):
    break
    
doc_paths = []

for i in range(0, len(files)):
    doc_paths.append(os.path.join(r'modules/docs2', files[i]))
    
for i in range(0, len(files_new)):
    doc_paths.append(os.path.join(r'modules/docs3', files_new[i]))

pay = ['Порядок оплаты', 'Условия оплаты','Условия   оплаты', 'Порядок и условия оплаты договорной цены', 'Порядок расчётов','Порядок расчетов', 'СТОИМОСТЬ УСЛУГ И ПОРЯДОК РАСЧЕТОВ', 'ПОРЯДОК И СРОКИ ОПЛАТЫ УСЛУГ', 'ПОРЯДОК ОПЛАТЫ', 'ПОРЯДОК РАСЧЕТОВ И СУММА ДОГОВОРА', 'ЦЕНА И ПОРЯДОК РАСЧЕТОВ', 'ЦЕНА КОНТРАКТА, ПОРЯДОК И СРОК ОПЛАТЫ', 'ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ', 'ЦЕНА ДОГОВОРА И УСЛОВИЯ ОПЛАТЫ', 'ЦЕНА ДОГОВОРА, ПРИЕМКА И ПОРЯДОК ОПЛАТЫ', 'ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЕТОВ', 'цена ДОГОВОРА и порядок расчетов', 'СТОИМОСТЬ УСЛУГИ И ПОРЯДОК РАСЧЁТОВ', 'СТОИМОСТЬ УСЛУГ И ПОРЯДОК РАСЧЕТОВ', 'СТОИМОСТЬ РАБОТ И ПОРЯДОК РАСЧЕТА', 'СТОИМОСТЬ ДОГОВОРАИ ПОРЯДОК РАСЧЁТОВ', 'СТОИМОСТЬ РАБОТ И ПОРЯДОК РАСЧЁТА',
      'Стоимость и порядок расчетов','Стоимость и порядок оплаты', 'Стоимость работ и порядок расчётов', 'ЦЕНА КОНТРАКТА И ПОРЯДОК РАСЧЕТОВ', 'РАЗМЕР ПЛАТЫ И ПОРЯДОК РАСЧЕТОВ', 'ПОРЯДОК ОСУЩЕСТВЛЕНИЯ ПЛАТЕЖЕЙ И РАСЧЕТОВ',' Стоимость услуг, порядок и сроки оплаты услуг', 'Порядок осуществления платежей и расчетов', 'Цена Контракта. Порядок и срок оплаты выполненных работ', 'Стоимость услуг и порядок расчетов','Стоимость услуг и порядок расчётов', 'Стоимость услуг и порядок оплаты', 'Стоимость Услуг и порядок оплаты','ЛИЦЕНЗИОННОЕ ВОЗНАГРАЖДЕНИЕ', 'УСЛОВИЯ РАСЧЕТА', 'ЦЕНЫ, ТАРИФЫ, СТОИМОСТЬ УСЛУГ И УСЛОВИЯ ОПЛАТЫ',
      'Цена Контракта и порядок расчетов','Цена контракта и порядок расчетов','ПОРЯДОК И СРОК ОПЛАТЫ РАБОТ', 'Порядок расчетов и платежей', 'ПРИЕМКА УСЛУГ и ПОРЯДОК РАСЧЕТОВ', 'СТОИМОСТЬ И ПОРЯДОК РАСЧЕТОВ ПО ДОГОВОРУ', 'СТОИМОСТЬ ТОВАРА И ПОРЯДОК РАСЧЕТОВ', 'раздел 3. порядок уплаты цены товара', 'Стоимость Работ и порядок оплаты.', 'Цена услуг по Договору', 'АРЕНДНАЯ ПЛАТА И ПОРЯДОК РАСЧЕТОВ ПО КОНТРАКТУ', 'СТОИМОСТЬ УСЛУГ И ПОРЯДОК ВЗАИМОРАСЧЕТОВ', 'УСЛОВИЯ И ПОРЯДОК  РАСЧЕТОВ', 'УСЛОВИЯ  И  ПОРЯДОК  РАСЧЕТОВ', 'Стоимость ДОГОВОРА и порядок расчетов', 'ЦЕНА ДОГОВОРА и ПОРЯДОК РАСЧЁТОВ', 'УСЛОВИЯ ОПЛАТЫ ВОЗНАГРАЖДЕНИЯ ЗА ПРЕДОСТАВЛЯЕМЫЕ ПРАВА НА ИСПОЛЬЗОВАНИЕ ПРОГРАММ ДЛЯ ЭВМ',
      'ОБЩАЯ ЦЕНА ДОГОВОРА. ПОРЯДОК РАСЧЕТОВ', 'ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЕТОВ', 'Цена настоящего Договора и порядок расчётов', 'Цена Договора и порядок расчетов', 'Стоимость договора и порядок оплаты','Цена договора и порядок расчетов','Цена договора и порядок расчетов.', 'ПОРЯДОК РАСЧЕТОВ ПО ДОГОВОРУ', 'Общая цена настоящего Договора и порядок расчётов', 'Стоимость договор и порядок расчетов', 'Стоимость услуг, порядок приемки услуг и порядок расчетов', 'Стоимость услуг связи и порядок расчетов', 'Стоимость Услуг, порядок расчетов', 'Форма и порядок расчетов за Услуги', 
       'ПОРЯДОК СДАЧИ-ПРИЕМКИ И СРОКИ ОПЛАТЫ РАБОТ', 'Цена работ и порядок расчетов', 'Платежи и расчеты', 'ОБЩАЯ ЦЕНА И ПОРЯДОК РАСЧЕТОВ', 'СТОИМОСТЬ ДОГОВОРА И ПОРЯДОК РАСЧЕТОВ','СТОИМОСТЬ ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ', 'Цена и порядок расчетов', 'Стоимость договора и порядок расчетов', 'Стоимость Договора и порядок расчетов', 'Цена Договора и порядок оплаты', 'РАЗМЕР ВОЗНАГРАЖДЕНИЯ И ПОРЯДОК РАСЧЕТОВ', 'Цена договора и порядок оплаты', 'Цена Контракта, порядок и сроки оплаты', 'Цена Договора, платежи и расчеты по Договору', 'Цена контракта, порядок и сроки оплаты товара, работ, услуг', 'Цена контракта, порядок и срок оплаты',
      'Цена Контракта и порядок оплаты','ОПЛАТА РАБОТ', 'Цена контракта и порядок оплаты', 'ПОРЯДОК РАСЧЕТОВ','Стоимость предоставления услуги и порядок расчетов', 'РАЗМЕР ВОЗНАГРАЖДЕНИЯ. ПОРЯДОК РАСЧЕТОВ', 'Цена контрактаи порядок расчётов','РАСЧЕТЫ МЕЖДУ СТОРОНАМИ', 'РАСЧЕТЫ СТОРОН', 'ПЛАТЕЖИ И РАСЧЕТЫ', 'Цена договора.', 'ЦЕНА КОНТРАКТА И ПОРЯДОК РАСЧЁТОВ', 'ЦЕНА ДОГОВОРА, ПОРЯДОК И ФОРМА РАСЧЕТОВ', 'Плата по Договору и порядок расчетов',
       'Цена контракта и порядок расчётов', 'Размер лицензионного вознаграждения и порядок расчетов', 'Цена Контракта, порядок и сроки оплаты услуг', 'Порядок и условия платежей', 'ЦЕНА ГОСУДАРСТВЕННОГО КОНТРАКТА', 'СТОИМОСТЬ РАБОТ И ПОРЯДОК РАСЧЕТОВ', 'ЦЕНА ДОГОВОРА И ПОРЯДОК ОПЛАТЫ', 'СТОИМОСТЬ УСЛУГ ПО ДОГОВОРУ И ПОРЯДОК РАСЧЕТОВ', 'РАСЧЁТЫ МЕЖДУ СТОРОНАМИ',
      'Стоимость услуги за эксплуатационно-техническое обслуживание\nсредств оповещения и связи ГО', 'Стоимость (цена) услуг и порядок оплаты', 'Стоимость работ и расчеты', 'Стоимость Услуги и порядок расчетов', 'Цена договора и порядок расчета', 'Цена Соглашения и порядок оплаты']
 
def first_word_finder(x,y):
    
    d = {'Контракт': [], 'Условия оплаты': [],'Cтарт-слово':[], 'Стоп-слово': []}
    issue_frame = pd.DataFrame(data=d)
    issues = []
    wrong = []
    wrong_total = []
    
    for count1, docdoc in enumerate(x): # цикл по договорам и индексам
        missing = 0 # счетсик пропущенных документов
        
        try:
            data = getText(docdoc)
        except:
            issues.append([1,count1, 'ERROR_Drip'])
            wrong_total.append(docdoc)
            continue
    
        for count2, indexer in enumerate(y):
            try:
                data.index(indexer) # нахождение точного литерального совпадения в тексте варианта из предложенных заголовков
            except Exception as e:
                missing += 1 # если нет - счетчик
            else:
                issues.append([data.index(indexer),count1, indexer]) # если не было исключений то добавляем номер строки название и индексер в фрейм
        if len(pay) == missing: # если счетчик ненайденных равен длине всех то ничего не нашлось вывод error
            wrong.append(docdoc)
            issues.append([1,count1, 'ERROR'])
        
            
    issue_sort = pd.DataFrame(sorted(issues,reverse = True, key=lambda issues: issues[1]), columns= ['string number', 'queue number', 'issue'])
    issue_sort = issue_sort.groupby(['queue number'], as_index = False)[['string number', 'issue']].max()
    
    return issue_sort, wrong, wrong_total
    
issue_sort, wrong, wrong_total = first_word_finder(doc_paths, pay)

files.extend(files_new)
issue_sort['docs'] = files
issue_sort['docs_dir'] = doc_paths

d = {'Контракт': [], 'Старт-слово':[], 'Стоп-слово': []}
stop_start_frame = pd.DataFrame(data=d)

def check_word(x,y):
    www = preprocessor(x)
    for index, element in enumerate(y):
        if bool(re.search(www+'(.*)', preprocessor(element[0]))):
            break
            
    return y[index][1], y[index+1][1], y[index][0], y[index+1][0]
    
param_list = list(map(lambda x,y :[x,y], issue_sort.issue.values, doc_paths))


def body_code(parm):
    d = {'Контракт': [], 'Старт-слово':[], 'Стоп-слово': []}
    stop_start_frame = pd.DataFrame(data=d)
    for counter in parm:
        data = docx.Document(counter[1])
        
        points = ['р', "Р", 'Раздел договора', 'ПУНКТ УР1', 'Раздел', 'punkt_dog', 'LB Gov style 1'] # внутренние стили для стиля 1
        style_bold = []
        styles1 = []
        style_rome = []
        styles2 = []
        styles3 = []
        styles7 = []
        bold_simple = []
        
        
        for i in range(0, len(data.paragraphs)):
            for j in range(0, len(data.paragraphs[i].runs)):
                if (data.paragraphs[i].runs[j].font.bold == True or data.paragraphs[i].runs[j].font.bold == None) and len(data.paragraphs[i].text)<160 \
        and (data.paragraphs[i].runs[j].font.underline == None or data.paragraphs[i].runs[j].font.underline == False):
                    if (bool(re.search('^(\d{1}\W{1}|раздел)', data.paragraphs[i].text)) or data.paragraphs[i].text.isupper() or bool(re.search('^[А-Я\s]+.*([а-я]|[А-Я])$', data.paragraphs[i].text)))\
                                            and (bool(re.search('[:]', data.paragraphs[i].text)) == False or data.paragraphs[i].text.isupper()) :
                        if bool(re.search('(\d{1}\W{1}\d{1}|\d{1}\)|№|\d{3,})', data.paragraphs[i].text)) == False:
                            style_bold.extend([[data.paragraphs[i].text, i]])
                            break
        
        
        for i in range(len(data.paragraphs)):
            if bool(re.search('^(?:IV|III|II|I|V)', data.paragraphs[i].text)):
                style_rome.extend([[data.paragraphs[i].text,i]])
        
        for i in range(len(data.paragraphs)):
            for j in points: # стиль 1
                if bool(data.paragraphs[i].style):
                    if data.paragraphs[i].style.name == j and len(data.paragraphs[i].text) < 100:
                        if bool(re.search('.', data.paragraphs[i].text)):
                            styles1.extend([[data.paragraphs[i].text, i]])
                    
        for i in range(len(data.paragraphs)):    
            if bool(data.paragraphs[i].style):
                if (data.paragraphs[i].style.name == 'Heading 1' or 'Normal') and len(data.paragraphs[i].text) < 70: # стиль 2
                    if bool(re.search('.', data.paragraphs[i].text)):
                        if bool(re.search('[\d]', data.paragraphs[i].text)) and bool(re.search('(\d{1}\W{1}\d{1}|\d{1}\)|№|\d{3,})', data.paragraphs[i].text)) == False: 
                            styles2.extend([[data.paragraphs[i].text,i]])
                    
        for i in range(len(data.paragraphs)):  
            if bool(data.paragraphs[i].style):
                if data.paragraphs[i].style.name == 'Heading 1' and len(data.paragraphs[i].text) < 100: ### стиль 3
                    if bool(re.search('.', data.paragraphs[i].text)):
                        styles3.extend([[data.paragraphs[i].text, i]])
                        
                        
        for i in range(0, len(data.paragraphs)):
            for j in range(0, len(data.paragraphs[i].runs)):
                if data.paragraphs[i].runs[j].font.bold == True and len(data.paragraphs[i].text) < 100:
                    if bool(re.search('.', data.paragraphs[i].text)) and bool(re.search('(\d{1}\W{1}\d{1}|\d{1}\)|№|\d{3,})', data.paragraphs[i].text)) == False: 
                        bold_simple.extend([[data.paragraphs[i].text, i]])
                        break
            
        style_bold_j = 'z'.join(list(map(lambda x: str(x[0]), style_bold)))          
        styles1_j = 'z'.join(list(map(lambda x: str(x[0]), styles1)))
        styles2_j = 'z'.join(list(map(lambda x: str(x[0]), styles2)))
        styles3_j = 'z'.join(list(map(lambda x: str(x[0]), styles3)))
        style_rome_j = 'z'.join(list(map(lambda x: str(x[0]), style_rome)))
        bold_simple_j = 'z'.join(list(map(lambda x: str(x[0]), bold_simple)))
        
        if bool(re.search(counter[0], style_rome_j)):
            start_par, end_par, f_name, e_name = check_word(counter[0], style_rome)
            new_row = {'Контракт':counter[1], 'Старт-слово': f_name, 'Стоп-слово': e_name, 'Number' : [start_par, end_par],
                      'Style': 'new_roman'}
            stop_start_frame = stop_start_frame.append(new_row, ignore_index=True)
        
        elif bool(re.search(counter[0], styles1_j)): 
            start_par, end_par, f_name, e_name = check_word(counter[0], styles1)
            new_row = {'Контракт':counter[1], 'Старт-слово': f_name, 'Стоп-слово': e_name, 'Number' : [start_par, end_par],
                      'Style': '1'}
            stop_start_frame = stop_start_frame.append(new_row, ignore_index=True)
        
        elif bool(re.search(counter[0], style_bold_j)): 
            start_par, end_par, f_name, e_name = check_word(counter[0], style_bold)
            new_row = {'Контракт':counter[1], 'Старт-слово': f_name, 'Стоп-слово': e_name, 'Number' : [start_par, end_par],
                      'Style': 'bold'}
            stop_start_frame = stop_start_frame.append(new_row, ignore_index=True)
        
        elif bool(re.search(counter[0], styles2_j)) :
            start_par, end_par, f_name, e_name = check_word(counter[0], styles2)
            new_row = {'Контракт':counter[1], 'Старт-слово': f_name, 'Стоп-слово': e_name, 'Number' : [start_par, end_par],
                      'Style': '2'}
            stop_start_frame = stop_start_frame.append(new_row, ignore_index=True)
        
        elif bool(re.search(counter[0], styles3_j)):
            start_par, end_par, f_name, e_name = check_word(counter[0], styles3)
            new_row = {'Контракт':counter[1], 'Старт-слово': f_name, 'Стоп-слово': e_name, 'Number' : [start_par, end_par],
                      'Style': '3'}
            stop_start_frame = stop_start_frame.append(new_row, ignore_index=True)
            
        elif bool(re.search(counter[0], bold_simple_j)):
            start_par, end_par, f_name, e_name = check_word(counter[0], bold_simple)
            new_row = {'Контракт':counter[1], 'Старт-слово': f_name, 'Стоп-слово': e_name, 'Number' : [start_par, end_par],
                      'Style': 'Simple_bold'}
            stop_start_frame = stop_start_frame.append(new_row, ignore_index=True)
        else:
            new_row = {'Контракт': counter[1], 'Старт-слово': counter[0], 'Стоп-слово': 'missed1', 'Number' : 'missed1', 'Style': 'missed'}
            stop_start_frame = stop_start_frame.append(new_row, ignore_index=True)
            
    return stop_start_frame   
    
ddd = body_code(param_list)

stop_start_frame = ddd[ddd['Number'] != 'missed1']

d = {'Контракт': [], 'Шаблон':[]}
pattern_frame = pd.DataFrame(data=d)

for index, row in stop_start_frame.iterrows():
    kwargs = {'start': row['Number'][0]+1, 'end': row['Number'][1]}
    text = getText(row['Контракт'], **kwargs)
    if any(c.isalpha() for c in text):
        
        new_row = {'Контракт':row['Контракт'], 'Шаблон': text}
        pattern_frame = pattern_frame.append(new_row, ignore_index= True)
    
    else:
        dts = getText(row['Контракт'])
        matches = [dts[m.end(0):j.start(0)] for m,j in zip(re.finditer(row['Старт-слово'], dts), re.finditer(row['Стоп-слово'], dts))]
        text = max(matches, key=len)
        new_row = {'Контракт':row['Контракт'], 'Шаблон': text}
        pattern_frame = pattern_frame.append(new_row, ignore_index= True)
        
        
###

tokenizer_sber = AutoTokenizer.from_pretrained('data/sbert_nlu_mt_ru')
model_sber = AutoModel.from_pretrained('data/sbert_nlu_mt_ru')

device = torch.device('cpu')
if torch.cuda.is_available():
    device = torch.device('cuda:0')


pattern_frame = pattern_frame[['Контракт', 'Шаблон']]

dataset_eval = Dataset.from_pandas(pattern_frame)

target_cols = ['fake']
MAX_LENGTH = 512
def preprocess_function(examples):
    result = tokenizer_sber(
        examples['Шаблон'],
        padding='max_length', max_length=MAX_LENGTH, truncation=True)
    
    if any(col in examples for col in target_cols):
        labels = list(zip(*[examples[col] for col in target_cols]))
        return {**result, "labels": labels}
    else:
        return result

dataset_eval = dataset_eval.map(preprocess_function, batched=True)


test_dataloader = torch.utils.data.DataLoader(dataset_eval, batch_size=dataset_eval.data.num_rows,
                                              collate_fn=transformers.default_data_collator, num_workers = 2)


class MeanPooling(nn.Module):
    def __init__(self):
        super(MeanPooling, self).__init__()
        
    def forward(self, last_hidden_state, attention_mask):
        input_mask_expanded = attention_mask.unsqueeze(-1).expand(last_hidden_state.size()).float()
        sum_embeddings = torch.sum(last_hidden_state * input_mask_expanded, 1)
        sum_mask = input_mask_expanded.sum(1)
        sum_mask = torch.clamp(sum_mask, min=1e-9)
        mean_embeddings = sum_embeddings / sum_mask
        return mean_embeddings


class ModelInfer(nn.Module):
    
    def __init__(self, bert: model_sber):
        super().__init__()
        self.cnn = nn.Sequential(nn.Conv2d(in_channels=1, out_channels= 128,
                                           kernel_size= (512,768),stride = 2, padding = 1),
                                nn.MaxPool2d(kernel_size=2, stride=2, padding = 1),
                                nn.Flatten(),
                                nn.Dropout(0.2),
                                nn.ReLU())
        self.bert = bert
        self.pool = MeanPooling()
        self.last = nn.Sequential(
            nn.Linear(bert.config.hidden_size, 256),
            nn.GELU(),
            nn.Linear(256, 3),
        nn.Softmax(dim=1))
        
    def forward(self, batch):
        bert_output = self.bert(
            input_ids=batch['input_ids'].to(device),
            attention_mask=batch['attention_mask'].to(device),
            token_type_ids=batch['token_type_ids'].to(device)
        )['last_hidden_state']
        
        bert_output = self.pool(bert_output, batch['attention_mask'])
        
        return self.last(bert_output).squeeze(1)

model = ModelInfer(model_sber).to(device)


model.load_state_dict(torch.load('data/empty-epoch=43-val_loss=0.6964.ckpt')["state_dict"]);


@torch.no_grad()
def validation_submit(model, valid_loader):
    #preds = torch.empty(0)
    for batch in valid_loader:
        for key in batch.keys():
            batch[key] = batch[key].to(device)
        pred = model(batch)
        pred = pred.argmax(dim = 1)
        #preds = torch.cat((preds, pred), 0)
        
    return pred


preds = validation_submit(model, test_dataloader)
preds_numpy = preds.detach().numpy()


data_out = {'Contract': pattern_frame['Контракт'], 'Text' : pattern_frame['Шаблон'], 'Status': preds_numpy}
out_df = pd.DataFrame(data_out)
out_df['Status'].replace([0,1,2], ['помесячно', 'постоплата', 'предоплата'], inplace=True);
today = str(date.today())
out_df.to_csv(f"modules/simple_output_{today}.csv", index = False)
print(out_df.values.tolist())



